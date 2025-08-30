import sys
import os
import random
import pandas as pd
import re
import sqlite3
import tempfile
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QPushButton, QListWidget, QLabel, 
                             QFileDialog, QMessageBox, QSplitter, QTextEdit,
                             QListWidgetItem, QCheckBox, QScrollArea, QFrame,
                             QLineEdit, QGroupBox, QSpinBox, QTableWidget,
                             QTableWidgetItem, QHeaderView, QTabWidget, QComboBox,
                             QDialog, QFormLayout, QDialogButtonBox)
from PyQt5.QtCore import Qt, QSize, QUrl
from PyQt5.QtGui import QFont, QIcon, QPixmap, QColor
from PyQt5.QtWebEngineWidgets import QWebEngineView
import PyPDF2
from docx import Document

# Set HighDPI scaling before creating QApplication
QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)
QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)

class KeywordManagerDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Keyword Manager")
        self.setGeometry(200, 200, 600, 500)
        self.db_connection = self.create_db_connection()
        self.initUI()
        
    def create_db_connection(self):
        """Create database connection with proper path handling"""
        # Ensure data directory exists
        data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')
        if not os.path.exists(data_dir):
            os.makedirs(data_dir)
        
        db_path = os.path.join(data_dir, 'cv_shuffler.db')
        return sqlite3.connect(db_path)
        
    def initUI(self):
        layout = QVBoxLayout(self)
        
        # Category selection
        form_layout = QFormLayout()
        self.category_combo = QComboBox()
        self.load_categories()
        form_layout.addRow("Category:", self.category_combo)
        
        self.keyword_input = QLineEdit()
        form_layout.addRow("New Keyword:", self.keyword_input)
        
        add_keyword_btn = QPushButton("Add Keyword")
        add_keyword_btn.clicked.connect(self.add_keyword)
        form_layout.addRow("", add_keyword_btn)
        
        layout.addLayout(form_layout)
        
        # Keywords list
        layout.addWidget(QLabel("Keywords in Selected Category:"))
        self.keywords_list = QListWidget()
        layout.addWidget(self.keywords_list)
        
        # Delete button
        delete_btn = QPushButton("Delete Selected Keyword")
        delete_btn.clicked.connect(self.delete_keyword)
        layout.addWidget(delete_btn)
        
        # Dialog buttons
        button_box = QDialogButtonBox(QDialogButtonBox.Ok)
        button_box.accepted.connect(self.accept)
        layout.addWidget(button_box)
        
        # Load keywords for initial category
        self.category_combo.currentIndexChanged.connect(self.load_keywords)
        self.load_keywords()
        
    def load_categories(self):
        cursor = self.db_connection.cursor()
        cursor.execute("SELECT id, name FROM job_categories ORDER BY name")
        categories = cursor.fetchall()
        
        self.category_combo.clear()
        for category_id, category_name in categories:
            self.category_combo.addItem(category_name, category_id)
    
    def load_keywords(self):
        self.keywords_list.clear()
        category_id = self.category_combo.currentData()
        
        if category_id:
            cursor = self.db_connection.cursor()
            cursor.execute("SELECT keyword FROM keywords WHERE category_id = ? ORDER BY keyword", (category_id,))
            keywords = cursor.fetchall()
            
            for (keyword,) in keywords:
                self.keywords_list.addItem(keyword)
    
    def add_keyword(self):
        keyword = self.keyword_input.text().strip()
        category_id = self.category_combo.currentData()
        
        if not keyword:
            QMessageBox.warning(self, "Input Error", "Please enter a keyword.")
            return
            
        try:
            cursor = self.db_connection.cursor()
            cursor.execute("INSERT OR IGNORE INTO keywords (keyword, category_id) VALUES (?, ?)", 
                          (keyword, category_id))
            self.db_connection.commit()
            
            self.keyword_input.clear()
            self.load_keywords()
            
        except Exception as e:
            QMessageBox.critical(self, "Database Error", f"Could not add keyword: {str(e)}")
    
    def delete_keyword(self):
        current_item = self.keywords_list.currentItem()
        if not current_item:
            return
            
        keyword = current_item.text()
        reply = QMessageBox.question(self, "Confirm Delete", 
                                    f"Are you sure you want to delete the keyword '{keyword}'?")
        
        if reply == QMessageBox.Yes:
            try:
                cursor = self.db_connection.cursor()
                cursor.execute("DELETE FROM keywords WHERE keyword = ?", (keyword,))
                self.db_connection.commit()
                self.load_keywords()
                
            except Exception as e:
                QMessageBox.critical(self, "Database Error", f"Could not delete keyword: {str(e)}")
    
    def closeEvent(self, event):
        self.db_connection.close()
        event.accept()

class CVShufflerApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.cv_files = []
        self.selected_candidates = []
        self.keyword_matches = {}
        self.temp_files = []  # To keep track of temporary files
        self.db_connection = self.create_db_connection()
        self.initUI()
        
    def create_db_connection(self):
        """Create database connection with proper path handling"""
        # Ensure data directory exists
        data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')
        if not os.path.exists(data_dir):
            os.makedirs(data_dir)
        
        db_path = os.path.join(data_dir, 'cv_shuffler.db')
        return sqlite3.connect(db_path)
        
    def initUI(self):
        self.setWindowTitle('CV Shuffler and Candidate Selector')
        self.setGeometry(100, 100, 1400, 900)
        
        # Central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Main layout
        main_layout = QHBoxLayout(central_widget)
        
        # Left panel for controls and CV list
        left_panel = QWidget()
        left_panel.setMaximumWidth(500)
        left_layout = QVBoxLayout(left_panel)
        
        # Title
        title_label = QLabel("CV Shuffler & Candidate Selector")
        title_label.setFont(QFont("Arial", 16, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        left_layout.addWidget(title_label)
        
        # Keyword section
        keyword_group = QGroupBox("Keyword Filtering")
        keyword_layout = QVBoxLayout(keyword_group)
        
        # Keyword management button
        self.manage_keywords_btn = QPushButton("Manage Keywords Database")
        self.manage_keywords_btn.clicked.connect(self.manage_keywords)
        keyword_layout.addWidget(self.manage_keywords_btn)
        
        # Category selection
        category_layout = QHBoxLayout()
        category_layout.addWidget(QLabel("Job Category:"))
        self.category_combo = QComboBox()
        self.load_categories()
        category_layout.addWidget(self.category_combo)
        keyword_layout.addLayout(category_layout)
        
        # Load keywords button
        self.load_keywords_btn = QPushButton("Load Keywords from Category")
        self.load_keywords_btn.clicked.connect(self.load_keywords_from_category)
        keyword_layout.addWidget(self.load_keywords_btn)
        
        # Custom keywords input
        keyword_input_layout = QHBoxLayout()
        keyword_input_layout.addWidget(QLabel("Custom Keywords:"))
        self.keyword_input = QLineEdit()
        self.keyword_input.setPlaceholderText("Enter comma-separated keywords")
        keyword_input_layout.addWidget(self.keyword_input)
        keyword_layout.addLayout(keyword_input_layout)
        
        # Threshold input
        threshold_layout = QHBoxLayout()
        threshold_layout.addWidget(QLabel("Match threshold:"))
        self.threshold_spin = QSpinBox()
        self.threshold_spin.setRange(1, 50)
        self.threshold_spin.setValue(5)
        threshold_layout.addWidget(self.threshold_spin)
        threshold_layout.addStretch()
        keyword_layout.addLayout(threshold_layout)
        
        # Case sensitivity
        self.case_sensitive_check = QCheckBox("Case sensitive matching")
        keyword_layout.addWidget(self.case_sensitive_check)
        
        # Apply keywords button
        self.apply_keywords_btn = QPushButton("Apply Keyword Filter")
        self.apply_keywords_btn.clicked.connect(self.apply_keyword_filter)
        self.apply_keywords_btn.setEnabled(False)
        keyword_layout.addWidget(self.apply_keywords_btn)
        
        # Auto-select button
        self.auto_select_btn = QPushButton("Auto-Select Matching CVs")
        self.auto_select_btn.clicked.connect(self.auto_select_matching)
        self.auto_select_btn.setEnabled(False)
        keyword_layout.addWidget(self.auto_select_btn)
        
        left_layout.addWidget(keyword_group)
        
        # Buttons
        btn_layout = QHBoxLayout()
        self.load_btn = QPushButton("Load CVs")
        self.load_btn.clicked.connect(self.load_cvs)
        self.shuffle_btn = QPushButton("Shuffle CVs")
        self.shuffle_btn.clicked.connect(self.shuffle_cvs)
        self.shuffle_btn.setEnabled(False)
        
        btn_layout.addWidget(self.load_btn)
        btn_layout.addWidget(self.shuffle_btn)
        left_layout.addLayout(btn_layout)
        
        # CV list with match counts
        left_layout.addWidget(QLabel("CV Files (Match Count):"))
        self.cv_list = QListWidget()
        self.cv_list.itemSelectionChanged.connect(self.show_cv_preview)
        left_layout.addWidget(self.cv_list)
        
        # Selection buttons
        select_btn_layout = QHBoxLayout()
        self.select_btn = QPushButton("Select Candidate")
        self.select_btn.clicked.connect(self.select_candidate)
        self.select_btn.setEnabled(False)
        self.deselect_btn = QPushButton("Remove Selection")
        self.deselect_btn.clicked.connect(self.deselect_candidate)
        self.deselect_btn.setEnabled(False)
        
        select_btn_layout.addWidget(self.select_btn)
        select_btn_layout.addWidget(self.deselect_btn)
        left_layout.addLayout(select_btn_layout)
        
        # Selected candidates
        left_layout.addWidget(QLabel("Selected Candidates:"))
        self.selected_list = QListWidget()
        left_layout.addWidget(self.selected_list)
        
        # Export button
        self.export_btn = QPushButton("Export Selected Candidates")
        self.export_btn.clicked.connect(self.export_selected)
        self.export_btn.setEnabled(False)
        left_layout.addWidget(self.export_btn)
        
        # Right panel with tabs for CV preview and keyword analysis
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        
        # Create tabs
        self.tabs = QTabWidget()
        
        # CV Preview tab - using QWebEngineView for proper rendering
        preview_tab = QWidget()
        preview_layout = QVBoxLayout(preview_tab)
        preview_layout.addWidget(QLabel("CV Preview (Original Format):"))
        self.preview_view = QWebEngineView()
        preview_layout.addWidget(self.preview_view)
        self.tabs.addTab(preview_tab, "CV Preview")
        
        # Text Content tab (for keyword extraction)
        text_tab = QWidget()
        text_layout = QVBoxLayout(text_tab)
        text_layout.addWidget(QLabel("Text Content (For Keyword Analysis):"))
        self.text_view = QTextEdit()
        self.text_view.setReadOnly(True)
        text_layout.addWidget(self.text_view)
        self.tabs.addTab(text_tab, "Text Content")
        
        # Keyword Analysis tab
        analysis_tab = QWidget()
        analysis_layout = QVBoxLayout(analysis_tab)
        analysis_layout.addWidget(QLabel("Keyword Analysis:"))
        self.keyword_table = QTableWidget()
        self.keyword_table.setColumnCount(3)
        self.keyword_table.setHorizontalHeaderLabels(["Keyword", "Count", "Context"])
        self.keyword_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        analysis_layout.addWidget(self.keyword_table)
        self.tabs.addTab(analysis_tab, "Keyword Analysis")
        
        right_layout.addWidget(self.tabs)
        
        # Add panels to main layout
        main_layout.addWidget(left_panel, 1)
        main_layout.addWidget(right_panel, 2)
        
        # Status bar
        self.statusBar().showMessage('Ready')
        
    def load_categories(self):
        cursor = self.db_connection.cursor()
        cursor.execute("SELECT id, name FROM job_categories ORDER BY name")
        categories = cursor.fetchall()
        
        self.category_combo.clear()
        self.category_combo.addItem("Select a category", None)
        for category_id, category_name in categories:
            self.category_combo.addItem(category_name, category_id)
    
    def load_keywords_from_category(self):
        category_id = self.category_combo.currentData()
        if not category_id:
            QMessageBox.warning(self, "Selection Error", "Please select a job category first.")
            return
            
        cursor = self.db_connection.cursor()
        cursor.execute("SELECT keyword FROM keywords WHERE category_id = ? ORDER BY keyword", (category_id,))
        keywords = cursor.fetchall()
        
        keyword_list = [keyword[0] for keyword in keywords]
        self.keyword_input.setText(", ".join(keyword_list))
        self.apply_keywords_btn.setEnabled(True)
        
        self.statusBar().showMessage(f"Loaded {len(keyword_list)} keywords from {self.category_combo.currentText()}")
    
    def manage_keywords(self):
        dialog = KeywordManagerDialog(self)
        dialog.exec_()
        # Refresh categories in case new ones were added
        self.load_categories()
        
    def load_cvs(self):
        options = QFileDialog.Options()
        files, _ = QFileDialog.getOpenFileNames(
            self, "Select CV Files", "", 
            "All Supported Files (*.pdf *.docx *.txt);;PDF Files (*.pdf);;Word Documents (*.docx);;Text Files (*.txt)",
            options=options
        )
        
        if files:
            self.cv_files = files
            self.update_cv_list()
            self.shuffle_btn.setEnabled(True)
            self.apply_keywords_btn.setEnabled(True)
            self.statusBar().showMessage(f"Loaded {len(files)} CVs")
            
    def update_cv_list(self):
        self.cv_list.clear()
        for file_path in self.cv_files:
            file_name = os.path.basename(file_path)
            item = QListWidgetItem(file_name)
            item.setData(Qt.UserRole, file_path)
            
            # Show match count if available
            if file_path in self.keyword_matches:
                match_count = sum(self.keyword_matches[file_path].values())
                item.setText(f"{file_name} ({match_count} matches)")
                # Color code based on match count
                if match_count >= 10:
                    item.setBackground(QColor(200, 255, 200))  # Light green for high matches
                elif match_count >= 5:
                    item.setBackground(QColor(255, 255, 200))  # Light yellow for medium matches
            
            self.cv_list.addItem(item)
            
    def shuffle_cvs(self):
        random.shuffle(self.cv_files)
        self.update_cv_list()
        self.statusBar().showMessage("CVs shuffled")
        
    def show_cv_preview(self):
        current_item = self.cv_list.currentItem()
        if current_item:
            file_path = current_item.data(Qt.UserRole)
            
            # Display the original document format
            if file_path.endswith('.pdf'):
                # For PDF files, display directly in the web view
                self.preview_view.setUrl(QUrl.fromLocalFile(file_path))
            elif file_path.endswith('.docx'):
                # For DOCX files, convert to HTML for display
                try:
                    doc = Document(file_path)
                    html_content = "<div style='font-family: Arial, sans-serif; padding: 20px;'>"
                    for para in doc.paragraphs:
                        if para.text.strip():  # Only add non-empty paragraphs
                            html_content += f"<p>{para.text}</p>"
                    html_content += "</div>"
                    
                    self.preview_view.setHtml(html_content)
                except Exception as e:
                    self.preview_view.setHtml(f"<h3>Error displaying DOCX file: {str(e)}</h3>")
            else:
                # For text files, display as HTML with monospace font
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    html_content = f"<pre style='font-family: monospace; padding: 20px;'>{content}</pre>"
                    self.preview_view.setHtml(html_content)
                except Exception as e:
                    self.preview_view.setHtml(f"<h3>Error displaying text file: {str(e)}</h3>")
            
            # Extract text for keyword analysis
            content = self.extract_text_from_cv(file_path)
            self.text_view.setPlainText(content)
            self.select_btn.setEnabled(True)
            self.deselect_btn.setEnabled(True)
            
            # Update keyword analysis tab if keywords were applied
            if file_path in self.keyword_matches:
                self.update_keyword_table(file_path, content)
    
    def extract_text_from_cv(self, file_path):
        text = ""
        try:
            if file_path.endswith('.pdf'):
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
            elif file_path.endswith('.docx'):
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            else:  # Assume text file
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
        except Exception as e:
            text = f"Error reading file: {str(e)}"
            
        return text if text else "No text could be extracted from this file."
    
    def apply_keyword_filter(self):
        keywords_text = self.keyword_input.text().strip()
        if not keywords_text:
            QMessageBox.warning(self, "No Keywords", "Please enter some keywords to filter by.")
            return
            
        # Parse keywords
        keywords = [k.strip() for k in keywords_text.split(',') if k.strip()]
        case_sensitive = self.case_sensitive_check.isChecked()
        threshold = self.threshold_spin.value()
        
        self.statusBar().showMessage(f"Applying {len(keywords)} keywords to {len(self.cv_files)} CVs...")
        
        # Process each CV for keyword matches
        self.keyword_matches = {}
        for file_path in self.cv_files:
            content = self.extract_text_from_cv(file_path)
            self.keyword_matches[file_path] = self.find_keyword_matches(content, keywords, case_sensitive)
        
        # Update the list to show match counts
        self.update_cv_list()
        
        # Count CVs that meet the threshold
        matching_cvs = [fp for fp in self.cv_files if sum(self.keyword_matches[fp].values()) >= threshold]
        self.statusBar().showMessage(f"Found {len(matching_cvs)} CVs with at least {threshold} keyword matches")
        self.auto_select_btn.setEnabled(True)
        
    def find_keyword_matches(self, content, keywords, case_sensitive):
        matches = {keyword: 0 for keyword in keywords}
        
        if not case_sensitive:
            content = content.lower()
            
        for keyword in keywords:
            search_term = keyword if case_sensitive else keyword.lower()
            # Use regex to find whole word matches only
            pattern = r'\b' + re.escape(search_term) + r'\b'
            matches[keyword] = len(re.findall(pattern, content))
            
        return matches
        
    def update_keyword_table(self, file_path, content):
        matches = self.keyword_matches[file_path]
        
        # Clear and setup table
        self.keyword_table.setRowCount(len(matches))
        
        # Populate table with keyword matches
        for row, (keyword, count) in enumerate(matches.items()):
            self.keyword_table.setItem(row, 0, QTableWidgetItem(keyword))
            self.keyword_table.setItem(row, 1, QTableWidgetItem(str(count)))
            
            # Find context for the keyword
            context = self.find_keyword_context(content, keyword, 
                                              self.case_sensitive_check.isChecked())
            self.keyword_table.setItem(row, 2, QTableWidgetItem(context))
            
        # Sort by count descending
        self.keyword_table.sortItems(1, Qt.DescendingOrder)
        
    def find_keyword_context(self, content, keyword, case_sensitive):
        if not case_sensitive:
            content = content.lower()
            keyword = keyword.lower()
            
        # Find the first occurrence of the keyword
        pos = content.find(keyword)
        if pos == -1:
            return "Not found"
            
        # Extract context around the keyword
        start = max(0, pos - 40)
        end = min(len(content), pos + len(keyword) + 40)
        context = content[start:end]
        
        # Highlight the keyword in the context
        if case_sensitive:
            context = context.replace(keyword, f"[{keyword}]")
        else:
            # Case insensitive replacement
            pattern = re.compile(re.escape(keyword), re.IGNORECASE)
            context = pattern.sub(r"[\g<0>]", context)
            
        return context
        
    def auto_select_matching(self):
        threshold = self.threshold_spin.value()
        
        # Clear previous selections
        self.selected_list.clear()
        self.selected_candidates = []
        
        # Select all CVs that meet the threshold
        for file_path in self.cv_files:
            if sum(self.keyword_matches[file_path].values()) >= threshold:
                file_name = os.path.basename(file_path)
                item = QListWidgetItem(file_name)
                item.setData(Qt.UserRole, file_path)
                self.selected_list.addItem(item)
                self.selected_candidates.append(file_path)
        
        if self.selected_candidates:
            self.export_btn.setEnabled(True)
            self.statusBar().showMessage(f"Auto-selected {len(self.selected_candidates)} candidates meeting the threshold")
        else:
            self.statusBar().showMessage("No candidates meet the threshold criteria")
    
    def select_candidate(self):
        current_item = self.cv_list.currentItem()
        if current_item:
            file_path = current_item.data(Qt.UserRole)
            file_name = os.path.basename(file_path)
            
            # Check if already selected
            for i in range(self.selected_list.count()):
                item = self.selected_list.item(i)
                if item.data(Qt.UserRole) == file_path:
                    QMessageBox.information(self, "Already Selected", 
                                           "This candidate has already been selected.")
                    return
            
            # Add to selected list
            item = QListWidgetItem(file_name)
            item.setData(Qt.UserRole, file_path)
            self.selected_list.addItem(item)
            self.selected_candidates.append(file_path)
            self.export_btn.setEnabled(True)
            self.statusBar().showMessage(f"Selected candidate: {file_name}")
            
    def deselect_candidate(self):
        current_item = self.cv_list.currentItem()
        if current_item:
            file_path = current_item.data(Qt.UserRole)
            file_name = os.path.basename(file_path)
            
            # Remove from selected list
            for i in range(self.selected_list.count()):
                item = self.selected_list.item(i)
                if item.data(Qt.UserRole) == file_path:
                    self.selected_list.takeItem(i)
                    self.selected_candidates.remove(file_path)
                    break
                    
            if not self.selected_list.count():
                self.export_btn.setEnabled(False)
                
            self.statusBar().showMessage(f"Removed candidate: {file_name}")
                
    def export_selected(self):
        if not self.selected_candidates:
            QMessageBox.warning(self, "No Selection", "No candidates have been selected.")
            return
            
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Selected Candidates", "", 
            "CSV Files (*.csv);;Text Files (*.txt)", options=options
        )
        
        if file_path:
            try:
                # Create a detailed report with keyword matches
                data = []
                for cv_path in self.selected_candidates:
                    file_name = os.path.basename(cv_path)
                    content_preview = self.extract_text_from_cv(cv_path)
                    
                    # Get keyword matches if available
                    keyword_info = ""
                    if cv_path in self.keyword_matches:
                        matches = self.keyword_matches[cv_path]
                        keyword_info = "; ".join([f"{k}:{v}" for k, v in matches.items() if v > 0])
                    
                    # Limit preview length
                    if len(content_preview) > 200:
                        content_preview = content_preview[:200] + "..."
                    
                    data.append({
                        "File Name": file_name,
                        "Path": cv_path,
                        "Keyword Matches": keyword_info,
                        "Preview": content_preview
                    })
                
                df = pd.DataFrame(data)
                
                if file_path.endswith('.csv'):
                    df.to_csv(file_path, index=False)
                else:
                    with open(file_path, 'w') as f:
                        f.write(df.to_string(index=False))
                    
                QMessageBox.information(self, "Export Successful", 
                                       f"Selected candidates exported to {file_path}")
                self.statusBar().showMessage(f"Exported {len(self.selected_candidates)} candidates to {file_path}")
                
            except Exception as e:
                QMessageBox.critical(self, "Export Error", f"Error exporting data: {str(e)}")
                self.statusBar().showMessage(f"Export error: {str(e)}")
    
    def closeEvent(self, event):
        # Clean up temporary files
        for temp_file in self.temp_files:
            try:
                os.unlink(temp_file)
            except:
                pass
        self.db_connection.close()
        event.accept()

def main():
    # Create application instance
    app = QApplication(sys.argv)
    
    # Check if database exists, create if not
    data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')
    db_path = os.path.join(data_dir, 'cv_shuffler.db')
    
    if not os.path.exists(db_path):
        # Create database if it doesn't exist
        if not os.path.exists(data_dir):
            os.makedirs(data_dir)
        
        # Run the setup script
        try:
            import subprocess
            subprocess.run([sys.executable, "setup_database.py"], check=True)
        except:
            # If setup script fails, show error
            QMessageBox.critical(None, "Database Error", 
                                "Could not create database. Please run setup_database.py manually.")
            return
    
    window = CVShufflerApp()
    window.show()
    
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()